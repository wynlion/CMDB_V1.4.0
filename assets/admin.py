from django.contrib import admin
from assets import models
# Register your models here.
from assets import asset_handler
from import_export.admin import ImportExportModelAdmin
from import_export import resources
from assets.models import ReportLists
from openpyxl import Workbook
from django.http import HttpResponse
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from io import StringIO
from datetime import date
from CMDB import settings
from django.utils.safestring import mark_safe
from django.utils.html import format_html
from django.urls import reverse
from django.contrib.admin import widgets


class NewAssetAdmin(admin.ModelAdmin):
    list_display = ['asset_type', 'sn', 'model', 'manufacturer', 'c_time', 'm_time', 'approved']
    list_filter = ['asset_type', 'manufacturer', 'c_time']
    search_fields = ('sn',)
    list_per_page = 10

    actions = ['approve_selected_assets']

    def approve_selected_assets(self, request, queryset):
        selected = request.POST.getlist('_selected_action')
        success_upline_number = 0
        for asset_id in selected:
            obj = asset_handler.ApproveAsset(request, asset_id)
            ret = obj.asset_upline()
            if ret:
                success_upline_number += 1

        self.message_user(request, "成功批准  %s  条新资产上线！" % success_upline_number)

    approve_selected_assets.short_description = '批准新资产上线'


class AssetAdmin(admin.ModelAdmin):
    list_display = ['asset_type', 'name', 'status', 'approved_by', 'c_time', 'm_time']
    list_filter = ['asset_type', 'status']
    list_per_page = 10


"""  已添加至projects App
class AllProjectsAdmin(admin.ModelAdmin):
    list_display = ['project_type', 'project_name', 'charge_person', 'members', 'project_status',
                    'project_status_percent', 'former_status_percent', 'start_day', 'expect_finished_day', 'memo']

    def members(self, obj):
        return [bt.name for bt in obj.group_members.all()]
    members.short_description = '项目成员'
    filter_horizontal = ['group_members']
    list_filter = ['project_type', 'project_status']
    list_per_page = 10
"""


class ReportResource(resources.ModelResource):

    class Meta:
        model = ReportLists
        fields = ('id', 'project_type', 'project_type', 'report_name', 'report_date', 'report_number')


class MemberAdmin(admin.ModelAdmin):
    list_display = ['name', 'belong_to_section', 'belong_to_group']
    list_filter = ['belong_to_section', 'belong_to_group']
    list_per_page = 10


class ReportListsAdmin(ImportExportModelAdmin):
    list_display = ['project_type', 'project_name', 'commissioned_units', 'report_name', 'report_date',
                    'report_number']
    search_fields = ['project_type', 'project_name', 'commissioned_units', 'report_name', 'report_date',
                     'report_number']
    list_filter = ['project_type', 'report_date']
    list_per_page = 10
    resource_class = ReportResource
    actions = ['print_selected_reports', 'export_selected_excel', 'export_to_word']

    def export_selected_excel(self, request, queryset):
        meta = self.model._meta  # 用于定义文件名, 格式为: app名.模型类名
        field_names = [field.name for field in meta.fields]  # 模型所有字段名

        response = HttpResponse(content_type='application/msexcel')  # 定义响应内容类型
        response['Content-Disposition'] = f'attachment; filename={meta}.xlsx'  # 定义响应数据格式
        wb = Workbook()  # 新建Workbook
        ws = wb.active  # 使用当前活动的Sheet表
        ws.append(field_names)  # 将模型字段名作为标题写入第一行
        for obj in queryset:  # 遍历选择的对象列表
            print(obj)
            for field in field_names:
                data = [f'{getattr(obj, field)}' for field in field_names]  # 将模型属性值的文本格式组成列表
            ws.append(data)  # 写入模型属性值
        wb.save(response)  # 将数据存入响应内容
        return response
    export_selected_excel.short_description = '导出到EXCEL'  # 该动作在admin中显示的文字

    # 功能未完善
    def export_to_word(self, request):
        document = Document()
        docx_title = "TEST_DOCUMENT.docx"
        # ---- Cover Letter ----
        # document.add_picture((r'%s/media/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
        document.add_paragraph()
        document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))
        document.add_paragraph('Dear Sir or Madam:')
        document.add_paragraph('We are pleased to help you with your widgets.')
        document.add_paragraph('Please feel free to contact me for any additional information.')
        document.add_paragraph('I look forward to assisting you in this project.')
        document.add_paragraph()
        document.add_paragraph('Best regards,')
        document.add_paragraph('Acme Specialist 1]')
        document.add_page_break()

        # Prepare document for download
        # -----------------------------
        f = StringIO()
        document.save(f)
        length = f.tell()
        f.seek(0)
        response = HttpResponse(
            f.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=' + docx_title
        response['Content-Length'] = length
        return response
    export_to_word.short_description = '导出到WORD'  # 该动作在admin中显示的文字


class ContractAdmin(admin.ModelAdmin):
    list_display = ['sn', 'name', 'price', 'start_day', 'end_day']
    list_per_page = 10


class EventLogAdmin(admin.ModelAdmin):
    list_display = ['name', 'asset', 'event_type', 'date', 'user']
    list_filter = ['event_type']
    list_per_page = 10


class IDCAdmin(admin.ModelAdmin):
    list_display = ['name', 'item', 'apparatus_status', 'apparatus_available']
    list_per_page = 10
    search_fields = ['name', 'item', 'apparatus_status', 'apparatus_available']


class LogFileAdmin(ImportExportModelAdmin):
    list_display = ['id', 'file', 'file_name', 'create_time', 'host_ip', 'memo']
    readonly_fields = ['headshot_image']

    def headshot_image(self, obj):
        return mark_safe('<img src="{url}" width="440px" height=275px />'.format(
            url=obj.headshot.url,
            width=obj.headshot.width,
            height=obj.headshot.height,
            )
        )

    list_per_page = 10
    search_fields = ['id', 'file_name', 'create_time', 'host_ip', 'memo']


class DownloadFileWidget(widgets.AdminFileWidget):
    id = None
    template_name = 'assets/download_file_input.html'

    def __init__(self, id, attrs=None):
        self.id = id
        super().__init__(attrs)

    def get_context(self, name, value, attrs):
        context = super().get_context(name, value, attrs)
        print(self, name, value, attrs, self.id)
        context['download_url'] = reverse('attachment', kwargs={'pk': self.id})
        return context


class AttachmentAdmin(admin.ModelAdmin):
    list_display = ['id', 'standard_type', 'box_type', 'number', 'name', 'standard_ID',
                    'memo', 'update_situation', '_get_download_url']
    search_fields = ['name', 'standard_ID']
    list_per_page = 10
    list_filter = ['standard_type']
    my_id_for_formfield = None

    def get_form(self, request, obj=None, **kwargs):
        if obj:
            self.my_id_for_formfield = obj.id
        return super(AttachmentAdmin, self).get_form(request, obj=obj, **kwargs)

    def formfield_for_dbfield(self, db_field, **kwargs):
        if self.my_id_for_formfield:
            if db_field.name == 'file':
                kwargs['widget'] = DownloadFileWidget(id=self.my_id_for_formfield)

        return super(AttachmentAdmin, self).formfield_for_dbfield(db_field, **kwargs)

    def _get_download_url(self, instance):
        return format_html('<a href="{}">{}</a>', reverse('attachment', kwargs={'pk': instance.id}), instance.name)

    _get_download_url.short_description = '下载文件'


admin.site.site_header = '广州交投检测用户系统'
admin.site.site_title = '广州交投检测用户系统'
admin.site.index_title = '资产管理'
admin.site.register(models.Asset, AssetAdmin)
admin.site.register(models.Server)
admin.site.register(models.StorageDevice)
admin.site.register(models.SecurityDevice)
admin.site.register(models.NetworkDevice)
admin.site.register(models.Software)
admin.site.register(models.BusinessUnit)
admin.site.register(models.Contract, ContractAdmin)
admin.site.register(models.Tag)
admin.site.register(models.IDC, IDCAdmin)
admin.site.register(models.Manufacturer)
admin.site.register(models.CPU)
admin.site.register(models.Disk)
admin.site.register(models.NIC)
admin.site.register(models.RAM)
admin.site.register(models.EventLog, EventLogAdmin)
admin.site.register(models.NewAssetApprovalZone, NewAssetAdmin)
# admin.site.register(models.AllProjects, AllProjectsAdmin)
admin.site.register(models.ReportLists, ReportListsAdmin)
admin.site.register(models.LogFile, LogFileAdmin)
admin.site.register(models.Member, MemberAdmin)
admin.site.register(models.Attachment, AttachmentAdmin)
# admin.site.register(models.Article, ArticleAdmin)
# admin.site.register(models.ReadNum, ReadNumAdmin)


"""新增测试功能"""

"""  待开发
class IssueAdmin(admin.ModelAdmin, ExportCsvMixin):
    fields = ('key', 'summary', 'status', 'project',
              'origin', 'components', 'prj_level', 'prj_category',
              'assignee', 'origin_person', 'pm', 'dev_manager', 'test_manager', 'tester', 'fe_dev', 'backend_dev',
              'plan_begin', 'plan_end', 'fe_plan_begin', 'fe_plan_end', 'test_plan_begin',
              'test_plan_end', 'backend_plan_begin', 'backend_plan_end',
              'created', 'reopen', 'prd_begin', 'prd_end', 'dev_begin', 'dev_end',
              'test_begin', 'test_end', 'pm_check', 'ready', 'pause', 'done',
              'pm_take', 'dev_take', 'test_take', 'total_take',
              'tags',
              )
    readonly_fields = fields
    list_display = ('key', 'summary', 'status', 'origin', 'components', 'created', 'visit')
    list_filter = ('origin', 'components', 'status', 'tags')
    search_fields = ('key', 'summary')
    date_hierarchy = 'created'
    actions = ['export_as_excel']
"""


"""  待测试
class ArticleAdmin(admin.ModelAdmin):
    list_display = ('id', 'title', 'author', 'text', 'readnum')
"""


"""  待测试
class ReadNumAdmin(admin.ModelAdmin):
    list_display = ('id', 'read_num_data', 'article')
"""


"""
class LogFileAdmin(admin.ModelAdmin, ExportCsvMixin):

    readonly_fields = [..., "headshot_image"]

    def headshot_image(self, obj):
        return mark_safe('<img src="{url}" width="{width}" height={height} />'.format(
            url = obj.headshot.url,
            width=obj.headshot.width,
            height=obj.headshot.height,
            )
    )  
"""

